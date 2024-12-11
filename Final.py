import gradio as gr
import pdfplumber
import docx
import os
import aiofiles
import asyncio
from sentence_transformers import SentenceTransformer, util
from transformers import pipeline
import logging
import nltk
from nltk.corpus import stopwords, wordnet
from nltk.tokenize import word_tokenize
from typing import List, Dict, Any, Optional
import torch
import time  # Added for time measurements
import re
from dateutil import parser as date_parser
from langchain_community.document_loaders import WebBaseLoader
from sklearn.cluster import KMeans
import numpy as np
from owlready2 import get_ontology, sync_reasoner_pellet
import random
import gym
from stable_baselines3 import PPO
import matplotlib.pyplot as plt
import pandas as pd
#from pronto import Ontology

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Initialize the sentence transformer model for re-ranking
sentence_transformer = SentenceTransformer('paraphrase-MiniLM-L6-v2', device='cpu')

# Initialize the summarization pipeline with dynamic device setting
device = 0 if torch.cuda.is_available() else -1
logging.info(f"Using device: {'GPU' if device == 0 else 'CPU'}")
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)

# Document class
class Document:
    def __init__(self, page_content: str, metadata: Dict[str, Any]):
        self.page_content = page_content
        self.metadata = metadata

# DocumentLoader class
class DocumentLoader:
    @staticmethod
    async def load_pdf(file_path: str) -> List[Document]:
        docs = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:  # Access pages correctly
                    text = page.extract_text()  # Extract text from each page
                    if text:  # Only add non-empty pages
                        docs.append(Document(page_content=text, metadata={"source": file_path}))
        except Exception as e:
            logging.error(f"Error loading PDF document: {e}")
        return docs

    @staticmethod
    async def load_docx(file_path: str) -> List[Document]:
        docs = []
        try:
            doc = await asyncio.to_thread(docx.Document, file_path)
            text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            docs.append(Document(page_content=text, metadata={"source": file_path}))
        except Exception as e:
            logging.error(f"Error loading DOCX document: {e}")
        return docs

    @staticmethod
    async def load_txt(file_path: str) -> List[Document]:
        docs = []
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            docs.append(Document(page_content=text, metadata={"source": file_path}))
        except Exception as e:
            logging.error(f"Error loading text document: {e}")
        return docs

    @staticmethod
    async def load_web(url: str) -> List[Document]:
        docs = []
        try:
            loader = WebBaseLoader(web_paths=(url,), bs_kwargs=dict())
            web_docs = loader.load()
            for doc in web_docs:
                docs.append(Document(page_content=doc.content, metadata={"source": url}))
        except Exception as e:
            logging.error(f"Error loading web document: {e}")
        return docs
    
    @staticmethod
    async def load_ann(file_path: str) -> List[Document]:
        docs = []
        try:
            # .ann files will be treated as text files
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
            docs.append(Document(page_content=text, metadata={"source": file_path}))
        except Exception as e:
            logging.error(f"Error loading ANN document: {e}")
        return docs

    @staticmethod
    async def load_local_docs(files) -> List[Document]:
        loader_map = {
            '.pdf': DocumentLoader.load_pdf,
            '.docx': DocumentLoader.load_docx,
            '.txt': DocumentLoader.load_txt,
            '.ann': DocumentLoader.load_ann
        }

        docs = []
        tasks = []

        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext == '.ann':
                # Rename .ann to .txt if needed
                new_file_path = file.replace('.ann', '.txt')
                os.rename(file, new_file_path)
                file = new_file_path
            
            loader = loader_map.get(ext)
            if loader:
                print(f"Loading file: {file}")  # Debug: Print each filename
                tasks.append(loader(file))
            else:
                logging.error(f"Unsupported file format: {file}")

        if tasks:
            docs_chunks = await asyncio.gather(*tasks)
            for chunk in docs_chunks:
                if chunk:  # Check if the chunk has content
                    print(f"Loaded {len(chunk)} documents from {chunk[0].metadata['source']}")  # Debug: Print source info
                docs.extend(chunk)
        
        print(f"Total documents loaded: {len(docs)}")  # Debug: Print total document count
        return docs

# HybridRetriever class
class HybridRetriever:
    def __init__(self):
        self.documents = []
        self.document_embeddings = None

    def add_documents(self, docs: List[Document]):
        self.documents.extend(docs)
        if docs:
            self.document_embeddings = sentence_transformer.encode(
                [doc.page_content for doc in self.documents],
                convert_to_tensor=True
            )

    def retrieve_docs(self, query: str, top_k=5) -> List[Document]:
        if self.document_embeddings is None or self.document_embeddings.numel() == 0:
            logging.error("No embeddings available to perform retrieval.")
            return []

        query_embedding = sentence_transformer.encode(query, convert_to_tensor=True)
        cosine_scores = util.pytorch_cos_sim(query_embedding, self.document_embeddings)[0]
        top_indices = cosine_scores.argsort(descending=True)[:top_k]
        return [self.documents[i] for i in top_indices]

    def cluster_documents(self, n_clusters: int = 5):
        if self.document_embeddings is None or self.document_embeddings.numel() == 0:
            logging.error("No embeddings available to perform clustering.")
            return []
        
        embeddings_np = self.document_embeddings.cpu().detach().numpy()
        kmeans = KMeans(n_clusters=n_clusters, random_state=0).fit(embeddings_np)
        clusters = {i: [] for i in range(n_clusters)}
        for idx, label in enumerate(kmeans.labels_):
            clusters[label].append(self.documents[idx])
        
        return clusters

# Query expansion function
def query_expansion(query: str) -> str:
    words = word_tokenize(query)
    expanded_query = []
    for word in words:
        synonyms = wordnet.synsets(word)
        expanded_query += [syn.lemmas()[0].name() for syn in synonyms][:2]
    return ' '.join(set(expanded_query + words))

# Temporal relation extraction
def extract_temporal_relations(text: str) -> str:
    temporal_keywords = ["before", "after", "next", "then", "subsequently", "previously", "during", "until", "when", "while", "first", "second", "finally"]
    date_patterns = [r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b", r"\b\d{4}-\d{2}-\d{2}\b"]

    temporal_relations = []
    sentences = nltk.sent_tokenize(text)
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in temporal_keywords):
            temporal_relations.append(f"[Temporal] {sentence}")
        for pattern in date_patterns:
            if re.search(pattern, sentence):
                temporal_relations.append(f"[Temporal] {sentence}")
                break

    return '\n'.join(temporal_relations)

# Causal relation extraction
def extract_causal_relations(text: str) -> str:
    causal_keywords = ["because", "due to", "therefore", "thus", "hence", "as a result", "leads to", "results in", "causes", "consequently"]

    causal_relations = []
    sentences = nltk.sent_tokenize(text)
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in causal_keywords):
            causal_relations.append(f"[Causal] {sentence}")

    return '\n'.join(causal_relations)

# Conditional relation extraction
def extract_conditional_relations(text: str) -> str:
    conditional_keywords = ["if", "unless", "provided that", "assuming that", "in case", "should"]

    conditional_relations = []
    sentences = nltk.sent_tokenize(text)
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in conditional_keywords):
            conditional_relations.append(f"[Conditional] {sentence}")

    return '\n'.join(conditional_relations)

# Comparative relation extraction
def extract_comparative_relations(text: str) -> str:
    comparative_keywords = ["greater", "less", "more", "fewer", "best", "worst", "better", "worse"]

    comparative_relations = []
    sentences = nltk.sent_tokenize(text)
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in comparative_keywords):
            comparative_relations.append(f"[Comparative] {sentence}")

    return '\n'.join(comparative_relations)

# Hierarchical relation extraction
def extract_hierarchical_relations(text: str) -> str:
    hierarchical_keywords = ["part of", "consists of", "includes", "composed of", "contains"]

    hierarchical_relations = []
    sentences = nltk.sent_tokenize(text)
    for sentence in sentences:
        if any(keyword in sentence.lower() for keyword in hierarchical_keywords):
            hierarchical_relations.append(f"[Hierarchical] {sentence}")

    return '\n'.join(hierarchical_relations)

# Summarization function for large documents
def batch_summarize_text(text: str) -> str:
    max_chunk_size = 1024
    sentences = text.split(". ")
    current_chunk = []
    all_summaries = []

    for sentence in sentences:
        current_chunk.append(sentence)
        if len(" ".join(current_chunk)) > max_chunk_size:
            summary = summarizer(" ".join(current_chunk), max_length=150, min_length=50, do_sample=False)
            all_summaries.append(summary[0]['summary_text'])
            current_chunk = []

    if current_chunk:
        summary = summarizer(" ".join(current_chunk), max_length=150, min_length=50, do_sample=False)
        all_summaries.append(summary[0]['summary_text'])

    return " ".join(all_summaries)

# Reinforcement Learning feedback mechanism
class FeedbackEnv(gym.Env):
    def __init__(self):
        super(FeedbackEnv, self).__init__()
        self.action_space = gym.spaces.Discrete(2)  # Actions: Improve response or keep the same
        self.observation_space = gym.spaces.Box(low=0, high=1, shape=(10,), dtype=np.float32)  # Placeholder state representation

    def reset(self):
        return np.zeros(10, dtype=np.float32)

    def step(self, action):
        reward = random.random() if action == 1 else 0  # Reward based on user feedback
        done = True  # Single-step environment for simplicity
        return np.zeros(10, dtype=np.float32), reward, done, {}

feedback_env = FeedbackEnv()
model = PPO('MlpPolicy', feedback_env, verbose=1)

# RetrieverManager class
class RetrieverManager:
    def __init__(self):
        self.hybrid_retriever = HybridRetriever()
        self.loading_time = 0
        self.query_time = 0

    async def load_and_retrieve_docs(self, files, url: Optional[str] = None) -> Optional[str]:
        start_time = time.time()  # Start timing
        if url:
            docs = await DocumentLoader.load_web(url)
        else:
            docs = await DocumentLoader.load_local_docs(files)
        
        if not docs:
            return None
        self.hybrid_retriever.add_documents(docs)
        end_time = time.time()  # End timing
        self.loading_time = end_time - start_time  # Calculate loading time
        return "Documents loaded successfully. You can now ask questions."

    def retrieve_docs(self, question: str) -> List[Document]:
        expanded_query = query_expansion(question)
        return self.hybrid_retriever.retrieve_docs(expanded_query)

    def cluster_documents(self, n_clusters: int = 5):
        return self.hybrid_retriever.cluster_documents(n_clusters)

# Format documents for display
def format_docs(docs: List[Document]) -> str:
    return "\n\n".join([f"Source: {doc.metadata['source']}\n\n{doc.page_content}" for doc in docs])

# ChatHistory class
class ChatHistory:
    def __init__(self):
        self.chat_history = []

    def add_user_message(self, message: str):
        self.chat_history.append({'role': 'user', 'content': message})

    def add_assistant_message(self, message: str):
        self.chat_history.append({'role': 'assistant', 'content': message})

    def get_transcript(self) -> str:
        return "\n".join([f"**{msg['role'].capitalize()}:** {msg['content']}" for msg in self.chat_history])

# Plotting function

def plot_dataset_with_query_results(documents, query):
    if not documents:
        logging.error("No documents available for plotting.")
        return None

    embeddings_np = sentence_transformer.encode([doc.page_content for doc in documents], convert_to_tensor=False)
    df = pd.DataFrame(embeddings_np[:, :2], columns=['Dim_1', 'Dim_2'])  # Use the first 2 dimensions for visualization
    df['text'] = [doc.page_content[:50] + "..." for doc in documents]  # Truncated content for display

    plt.figure(figsize=(10, 6))
    plt.scatter(df['Dim_1'], df['Dim_2'], label='All Documents', color='blue', alpha=0.5)

    # Retrieve and mark relevant documents
    retrieved_docs = retriever_manager.retrieve_docs(query)
    if retrieved_docs:
        relevant_indices = [documents.index(doc) for doc in retrieved_docs]
        plt.scatter(df.loc[relevant_indices, 'Dim_1'], df.loc[relevant_indices, 'Dim_2'], color='red', label='Relevant Documents', marker='x', s=100)

    plt.title('Document Embeddings with Query Results')
    plt.xlabel('Dimension 1')
    plt.ylabel('Dimension 2')
    plt.legend()
    plt.grid(True)

    plt.savefig('/tmp/dataset_plot.png')
    plt.close()
    return '/tmp/dataset_plot.png'

# Function to handle questions and plot results
def handle_question_with_plot(question: str):
    response = rag_chain(question, retriever_manager, chat_history)
    query_time = retriever_manager.query_time
    plot_path = plot_dataset_with_query_results(retriever_manager.hybrid_retriever.documents, question)
    if plot_path:
        return f"{response}\n\n**Query Response Time:** {query_time:.2f} seconds", plot_path, ""
    else:
        return f"{response}\n\n**Query Response Time:** {query_time:.2f} seconds", None, ""


# RAG chain function
def rag_chain(question: str, retriever_manager: RetrieverManager, chat_history: ChatHistory) -> str:
    if not retriever_manager.hybrid_retriever.documents:
        return "No document loaded. Please load a document first."

    start_time = time.time()  # Start timing
    retrieved_docs = retriever_manager.retrieve_docs(question)
    if not retrieved_docs:
        retriever_manager.query_time = time.time() - start_time  # Record query time
        return "No relevant documents found for the question."

    formatted_context = format_docs(retrieved_docs)

    temporal_info = extract_temporal_relations(formatted_context)
    causal_info = extract_causal_relations(formatted_context)
    conditional_info = extract_conditional_relations(formatted_context)
    comparative_info = extract_comparative_relations(formatted_context)
    hierarchical_info = extract_hierarchical_relations(formatted_context)

    if "summarize" in question.lower():
        summarized_content = batch_summarize_text(formatted_context)
        response = f"**Summary of the document:**\n{summarized_content}\n\n**Temporal Relations:**\n{temporal_info}\n\n**Causal Relations:**\n{causal_info}\n\n**Conditional Relations:**\n{conditional_info}\n\n**Comparative Relations:**\n{comparative_info}\n\n**Hierarchical Relations:**\n{hierarchical_info}"
    else:
                response = f"**Response based on the retrieved documents:**\n{formatted_context}\n\n**Temporal Relations:**\n{temporal_info}\n\n**Causal Relations:**\n{causal_info}\n\n**Conditional Relations:**\n{conditional_info}\n\n**Comparative Relations:**\n{comparative_info}\n\n**Hierarchical Relations:**\n{hierarchical_info}"

    retriever_manager.query_time = time.time() - start_time  # Record query time

    chat_history.add_user_message(question)
    chat_history.add_assistant_message(response)

    return chat_history.get_transcript()

# Remove load button functionality, automatically load when file is uploaded
async def auto_load_docs(files, url=None):
    status = await retriever_manager.load_and_retrieve_docs(files, url)
    if status:
        loading_time = retriever_manager.loading_time
        return f"{status}\n\n**Loading Time:** {loading_time:.2f} seconds"
    else:
        return "Failed to load documents."

def auto_load_docs_sync(files, url=None):
    return asyncio.run(auto_load_docs(files, url))

def clear_data():
    retriever_manager.__init__()  # Reinitialize retriever
    chat_history.__init__()  # Reinitialize chat history
    return "All data cleared."

def handle_question(question: str):
    response = rag_chain(question, retriever_manager, chat_history)
    query_time = retriever_manager.query_time
    # Include the query time in the response
    full_response = f"{response}\n\n**Query Response Time:** {query_time:.2f} seconds"
    return full_response, ""

retriever_manager = RetrieverManager()
chat_history = ChatHistory()
# Set up the interface
with gr.Blocks() as iface:
    gr.Markdown("# Enhanced RAG System with Data Plotting", elem_id="header")

    # Document Upload Section
    file_input = gr.File(label="Upload Documents (PDF, DOCX, TXT)", file_count="multiple", type="filepath")
    url_input = gr.Textbox(label="Web URL", placeholder="Enter a URL to load web content...")
    loading_status = gr.Markdown()
    
    load_button = gr.Button("Load Document")
    load_button.click(fn=auto_load_docs_sync, inputs=[file_input, url_input], outputs=loading_status)

    # Clear Data Button
    clear_data_button = gr.Button("Clear All Data", variant="secondary")
    clear_data_output = gr.Markdown()
    clear_data_button.click(fn=clear_data, outputs=clear_data_output)

    # Chat Section
    chat_output = gr.Markdown(elem_id="chat-output")
    question_input = gr.Textbox(placeholder="Enter your question here...", lines=1)
    plot_output = gr.Image(type="filepath")  # Ensure this matches the file path format
    submit_button = gr.Button("Submit", variant="primary")
    submit_button.click(fn=handle_question_with_plot, inputs=[question_input], outputs=[chat_output, plot_output, question_input])

    # Custom Prompt for Training
    training_prompt = gr.Textbox(label="Custom Training Prompt", placeholder="Enter a custom prompt to train the model...")
    train_button = gr.Button("Train Model")
    train_output = gr.Markdown()
    
    def train_model(files, prompt):
        # Placeholder function for training the model using uploaded files and a custom prompt
        return f"Model trained with custom prompt: {prompt}"
    
    train_button.click(fn=train_model, inputs=[file_input, training_prompt], outputs=train_output)

    # Document Clustering Section
    cluster_output = gr.Markdown()
    cluster_button = gr.Button("Cluster Documents")
    
    def handle_clustering():
        clusters = retriever_manager.cluster_documents(n_clusters=5)
        cluster_text = ""
        for cluster_id, docs in clusters.items():
            cluster_text += f"**Cluster {cluster_id}:**\n"
            cluster_text += "\n\n".join([f"Source: {doc.metadata['source']}\n{doc.page_content[:200]}..." for doc in docs])
            cluster_text += "\n\n"
        return cluster_text
    
    cluster_button.click(fn=handle_clustering, outputs=cluster_output)

    # Layout organization
    with gr.Row():
        with gr.Column(scale=1):
            file_input
            url_input
            load_button
            loading_status
            clear_data_button
            clear_data_output

        with gr.Column(scale=2):
            chat_output
            question_input
            submit_button
            plot_output
            training_prompt
            train_button
            train_output
            cluster_button
            cluster_output

iface.launch()